from urllib import request
from flask import Flask,url_for
from flask import render_template,request, make_response
import sympy
from sympy.abc import *
from sympy import * 
from sympy.integrals.manualintegrate import integral_steps
import re
import decimal
import numpy as np
import time
import os
from scipy.integrate import quad
import base64
import tempfile
from PIL import Image
import io
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_TAB_ALIGNMENT

app = Flask(__name__)
@app.context_processor
def override_url_for():
    return dict(url_for=dated_url_for)

def dated_url_for(endpoint, **values):
    if endpoint == 'static':
        filename = values.get('filename', None)
        if filename:
            file_path = os.path.join(app.root_path,
                                     endpoint, filename)
            values['q'] = int(os.stat(file_path).st_mtime)
    return url_for(endpoint, **values)
@app.route("/download",methods=['GET','POST'])
def downloadzip():
    aa=request.form["file"]
    aa=aa[1:]
    aa=aa.replace('$','\\')
    aa=aa.split('#')
    
    document = Document()
    for i in range(len(aa)):
        if i%2==0:
            p=document.add_paragraph()
            p.add_run(text=aa[i], style=None).font.math = True
            tabstops = p.paragraph_format.tab_stops
            tabstops.add_tab_stop(0, WD_TAB_ALIGNMENT.START)
        else:
            p=document.add_paragraph()
            p.add_run(text=aa[i], style=None).font.math = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    filepath = os.path.join(tempfile.mkdtemp(),'test.docx')
    document.save(filepath)
    response = make_response()
    response.data  = open(filepath, "rb").read()
    response.headers['Content-Type'] = 'application/octet-stream'
    response.headers['Content-Disposition'] = 'attachment; filename=test.docx'
    return response

@app.route("/",methods=['GET','POST'])
def hello_world():
    aa=''
    bb='x'
    ur="location.href='./'"
    ur2="./"
    check2='none'
    img_data3=''
    img_data4=''
    if request.method=='POST':
        check2=request.form.get('hage')
        if (check2=='yes')or(check2=='no'):
            ur="location.href='./step'"
            ur2="./step"
            check2='yes'            
        aa,bb,mo,gr,rr=requ(check2)
        img_string=graph(simplify(rr),-2,2,mo)
        img_string2=graph(simplify(rr),-10,10,mo)
        if gr[0]==1:
            img_data3=graph(gr[3],-2,2,mo)
            img_data4=graph(gr[3],-10,10,mo)
        fir=gr[1]        
        res=gr[2]
    else:
        img_string=graph(x,-2,2,x)
        img_string2=graph(x**2,-2,2,x)
        fir=''
        res=''
    cc=''
    for i in range(len(aa)):
        cc=cc+'#'+aa[i]
    cc=cc.replace('\\','$')
    return render_template('hello.html',aa=aa,bb=bb,cc=cc,fir=fir,res=res,img_data4=img_data4,img_data3=img_data3,check2=check2,ur=ur,ur2=ur2,img_data=img_string,img_data2=img_string2,form=latex(eval(bb)))

@app.route("/step",methods=['GET','POST'])
def hello_world2():
    aa=''
    bb='x'
    ur="location.href='./step'"
    ur2="./step"
    check2='yes'
    img_data3=''
    img_data4=''
    if request.method=='POST':
        check2=request.form.get('hage')
        if check2=='yes':
            ur="location.href='./'"
            ur2="./"
            check2='none'
        else:
            check2='yes'
        aa,bb,mo,gr,rr=requ(check2)
        img_string=graph(simplify(rr),-2,2,mo)
        img_string2=graph(simplify(rr),-10,10,mo)
        if gr[0]==1:
            img_data3=graph(gr[3],-2,2,mo)
            img_data4=graph(gr[3],-10,10,mo)
        fir=gr[1]        
        res=gr[2]
    else:
        img_string=graph(x,-2,2,x)
        img_string2=graph(x**2,-2,2,x)
        res=''
        fir=''
    cc=''
    for i in range(len(aa)):
        cc=cc+'#'+aa[i]
    cc=cc.replace('\\','$')
    return render_template('hello.html',aa=aa,bb=bb,cc=cc,fir=fir,res=res,img_data4=img_data4,img_data3=img_data3,check2=check2,ur=ur,ur2=ur2,img_data=img_string,img_data2=img_string2,form=latex(eval(bb)))



@app.route("/連立方程式",methods=['GET','POST'])
def renritu():
    bb=''
    if request.method=='POST':
        try:
            
            x = Symbol(request.form["username"])
            y = Symbol(request.form["username1"])
            
            equation1 = simplify(request.form["username2"])
            equation2 = simplify(request.form["username3"])
            
            aa=[latex(solve([equation1, equation2]))]
        except:
            aa=['error']
    return render_template('hello.html',aa=aa,bb=bb)


@app.route("/微分方程式",methods=['GET','POST'])
def bibunhou():
    bb=''
    if request.method=='POST':
        '''
        try:
            bb=request.form["username3"]
            y = Function('y')(x)
            eq =simplify(bb)
            aa=[latex(eq)+"=0,",latex(dsolve(eq, ics={y.subs(x,0):request.form["username4"],diff(y,x,1).subs(x,request.form["username5"]):1}))]
            aa.append(latex(dsolve(eq)))
            
            
        except:
            aa=['error']
        '''

        m=[[ 2, 2,-1, 1],
     [-1, 1, 2, 1],
     [ 1, 2, 0, 1],
     [ 2, 0, 1, 1]]
        lat=[]
        aa=det(m,lat)
        bb=''


    return render_template('hello.html',aa=aa,bb=bb)

@app.route("/book",methods=['GET','POST'])
def book():
    return render_template('book.html')

def graph_plot(form,mini,maxi,moji):
    f = lambdify(moji, form, 'numpy')
    max=30
    xa=np.arange(mini, maxi, (maxi-mini)/1000)
    for k in range(100):
        li=0
        ya=[]
        xaa=[]
        yaa=[]
        hyoji=0
        for i in range(len(xa)):
            if abs(f(xa[i]))>max:
                if abs(li-i)>1:
                    xaa.append(xa[li+1:i])
                    yaa.append(ya[li+1:i])
                    hyoji=hyoji+i-li
                li=i
                ya.append(f(xa[i]))
            else:
                ya.append(f(xa[i]))
        hyoji=hyoji+i-li
        if hyoji/len(xa)>0.6:
            break
        max=max*2
    xaa.append(xa[li+1:i])
    yaa.append(ya[li+1:i])
    return xaa,yaa

def graph(form,min,max,moji):
    try:
        xaa,yaa=graph_plot(form,min,max,moji)
        fig = plt.figure()
        plt.title(str(min)+'<'+str(moji)+'<'+str(max))
        plt.xlabel(str(moji))
        plt.grid()
        for i in range(len(xaa)):
            plt.plot(xaa[i],yaa[i],color='blue')
    except:
        fig = plt.figure()
        plt.grid()    
    plt.gca().spines['right'].set_visible(False)
    plt.gca().spines['top'].set_visible(False)
    filepath = tempfile.mkdtemp()+"/a.png"
    #filepath = "{}/".format(tempfile.gettempdir()) + datetime.now().strftime("%Y%m%d%H%M%S") + ".png"
    fig.savefig(filepath)
    image = Image.open(filepath)
    image = np.asarray(image)     
    filtered = Image.fromarray(image)
    # base64でエンコード
    buffer = io.BytesIO()
    filtered.save(buffer, format="PNG")
    img_string = base64.b64encode(buffer.getvalue()).decode().replace("'", "")
    return img_string

def requ(check2):
    gr=[0,'',0,0] #　グラフの表示の有無、ボタンに表示する関数、入力解釈、プロットする関数
    ti1=time.time()
    bb=request.form["username"]
    bb=siki(bb)
    aa=[]

    try:
        if (bb!=bb.replace('\\','')):
            form=mat(bb)
        else:
            form=simplify(bb)
    except:
        bb='x'
        form=simplify('x')

    moj,rr=parse(form)
    mo=simplify(moj[0])
    try:
        aa=[request.form["tse"]]
        ini=request.form["ini"]
        end=request.form["end"]
        gr=[1,'\\int_{'+ini+'}^{'+end+'}('+latex(form)+')d'+latex(mo),'\\int{'+latex(form)+'}d'+latex(mo),0]
        if check2=='yes':
            aa=set_step(form,ini,end,1,[],mo)  
        else:
            x=Symbol('x')
            ans=integrate(form,(x,simplify(ini),simplify(end)))
            if (str(ans)!=str(ans).replace('Integ',''))and(str(ans)==str(ans).replace('Piecewise','')):
                f = lambdify(x, form, 'numpy')
                y = lambda x: f(x)
                if ini=='oo':
                    ini=10**10
                if end=='oo':
                    end=10**10
                if ini=='-oo':
                    ini=-10**10
                if end=='oo':
                    end=-10**10                
                ans=quad(y, int(ini), int(end))[0]
            aa=['','\\int_{'+ini+'}^{'+end+'}('+latex(form)+')dx='+latex(ans)]
    except KeyError:
         aa=aa
    if len(aa)==0:            
        try:
            if check2=='yes':
                aa=[request.form["se"]]
                aa=set_step(form,0,0,0,[],'x')
                gr=[1,'\\int{'+latex(form)+'}d'+latex(mo),'\\int{'+latex(form)+'}d'+latex(mo),0]
            else:
                aa=[request.form["se"]]
                x=Symbol('x')
                res=integrate(simplify(form), x)
                if (str(res).replace('Integ','')!=str(res)):
                    aa=['解析的な結果が見つかりません.',latex(Integral(form,x))+'=?']
                else:
                    aa=['','\\int('+latex(form)+')dx='+latex(res)]
                    gr=[1,'\\int{'+latex(form)+'}d'+latex(mo),'\\int{'+latex(form)+'}d'+latex(mo),res]
            
        except KeyError:
            aa=aa
        try:
            aa=[request.form["bi"]]
            aa=dif(form,[])
            if (len(aa)>3):
                aa[-1]=aa[-1]+'以上より'
                aa.append("("+latex(form)+")'="+latex(diff(form,x))+'.')
        except KeyError:
            aa=aa
        try:
            aa=[request.form["max"]]
            aa=maxi(form,check2)
        except KeyError:
            aa=aa
        try:
            aa=[request.form["sol"]]
            aa=equa(form,[])[2]
        except KeyError:
            aa=aa 
        try:
            aa=[request.form["series"]]
            x=Symbol('x')
            aa=['',latex(eval(bb))+'='+latex(series(simplify(bb),mo))]
        except KeyError:
            aa=aa
        try:
            aa=[request.form["lim"]]
            end=request.form["end2"]
            moji=request.form["moji"]
            moji=simplify(moji)
            x=Symbol('x')
            if (check2=='yes'):
                aa=limi(form,end,1,[''],moji)[1]
            else:
                aa=limi(form,end,0,[''],moji)[1]
            gr=[0,'\\lim_{'+latex(moji)+'\\to '+latex(end)+'}\\left('+latex(form)+'\\right)',0,0]
        except KeyError:
            aa=aa     
        try:
            aa=[request.form["lap"]]
            if check2=='yes':
                aa=lap(form,1)
            else:
                aa=lap(form,0)
            gr=[1,'\\mathcal{L}_{x} \\left['+latex(form)+'\\right](s)','\\mathcal{L}_{x} \\left['+latex(form)+'\\right](s)',0]
        except KeyError:
            aa=aa   
        try:
            aa=[request.form["inlap"]]
            if check2=='yes':
                aa=inlap(bb,1)
            else:
                aa=inlap(bb,0)
            gr=[1,'\\mathcal{L}_{s}^{-1} \\left['+latex(form)+'\\right](x)','\\mathcal{L}_{s}^{-1} \\left['+latex(form)+'\\right](x)',0]
        except KeyError:
            aa=aa
    ti2=time.time()      
    gr[1]='処理時間:'+str(int((ti2-ti1)*100)/100)+'s　　入力解釈:　'+gr[1]+'　を求める.'  
    return aa,bb,mo,gr,rr


def siki(form):
    kansu=['sqrt','$0###','abs','$1###','delta','$2###','asinh','$3###','acosh','$4###','atanh',
    '$5###','log','$6###','exp','$7###','Si','$8###','Ci','$9###','asin','1##','acos','2##','atan',
    '$3##','arcsinh','$3###','arccosh','$4###','arctanh','$5###','Ei','$4##','erfi','$5##','erf','$6##',
    'csc','$7##','sec','$8##','cot','$9##','arcsin','$1##','arccos','$2##','arctan','$3##','sinh','$0#',
    'cosh','$1#','tanh','$2#','sin','$3#','cos','$4#','tan','$5#','Ai','$6#','Bi','$7#']
    form=form.replace('^','**')
    form=form.replace(')(',')*(')
    for i in range(int(len(kansu)/2)):
        form=form.replace(kansu[2*i],kansu[2*i+1])
    if (form!=form.replace('#','')):
        print(form)
        for i in range(100):
            if i==len(form)-1:
                break
            if (form[i]=='#')and(form[i+1]!='(')and(form[i+1]!='#'):
                if i+2==len(form):
                    form=form[:i+1]+'('+form[i+1]+')'
                else:
                    for j in range(30):
                        if i+2+j==len(form):
                            form=form[:i+1]+'('+form[i+1:i+2+j]+')'
                            break
                        elif (form[i+2+j]=='$')or(form[i+2+j]=='+')or(form[i+2+j]=='-')or(form[i+2+j]=='*')or(form[i+2+j]=='/'):
                            form=form[:i+1]+'('+form[i+1:i+2+j]+')'+form[i+2+j:]
                            break
                            
    for i in range(10):
        form=form.replace(str(i)+'(',str(i)+'*(')
        form=form.replace(str(i)+'$',str(i)+'*$')
        form=form.replace(')'+str(i),')*'+str(i))
    alp=['a','b','c','d','e','f','g','h','i','j','k','l','m','n','o','p','q','r','s','t','u','v','w','x','y','z','A','B','C','D','E','F','G','H','I','J','K','L','M','N','O','P','Q','R','S','T','U','V','W','X','Y','Z']
    for i in alp:
        form=form.replace(i+'(',i+'*(')
        form=form.replace(i+'$',i+'*$')
        form=form.replace(')'+i,')*'+i)
        for j in range(10):
            form=form.replace(str(j)+i,str(j)+'*'+i)
        for j in alp:
            form=form.replace(j+i,j+'*'+i)
    form=form.replace(')$',')*$')
    for i in range(int(len(kansu)/2)):
        form=form.replace(kansu[2*i+1],kansu[2*i])
    return form

def steps(rule,lat,yobi,moji):
        rule_kind=str(rule)[0:15]
        if (rule_kind.replace('del','')!=rule_kind):
            rule.pop(0)
        if ((rule_kind.replace('ExpRule','')!=rule_kind)or(rule_kind.replace('PowerRule','')!=rule_kind)or(rule_kind.replace('TrigRule','')!=rule_kind)or(rule_kind.replace('LiRule','')!=rule_kind)or(rule_kind.replace('CiRule','')!=rule_kind)or(rule_kind.replace('SiRule','')!=rule_kind)or(rule_kind.replace('EiRule','')!=rule_kind)):
                if (len(lat)==0):
                    lat.append('積分の公式より')
                yobi.append(rule[2])
                lat.append(latex(Integral(rule[2],rule[3]))+'='+latex(integrate(rule[2],rule[3])))
                lat.append('となる. ')

        elif ((rule_kind.replace('ReciprocalRule','')!=rule_kind)or(rule_kind.replace('PiecewiseRule','')!=rule_kind)):
                if (len(lat)==0):
                    lat.append('積分の公式より')
                yobi.append(rule[1])
                lat.append(latex(Integral(rule[1],rule[2]))+'='+latex(integrate(rule[1],rule[2])))
                lat.append('となる. ')
        elif (rule_kind.replace('ConstantTimes','')!=rule_kind):
                if (len(lat)==0):
                    lat.append('積分の公式より')
                if check_rule(str(rule[2])[0:15])==1:
                    yobi.append(rule[0]*rule[1])
                    lat.append(latex(Integral(rule[0]*rule[1],rule[4]))+'='+latex(integrate(rule[0]*rule[1],rule[4])))
                    lat.append('となる. ')
                else:
                    lat=steps(rule[2],lat,yobi,moji)[0]
        elif ((rule_kind.replace('ErfRule','')!=rule_kind)or(rule_kind.replace('Fresnel','')!=rule_kind)):
                if (len(lat)==0):
                    lat.append('積分の公式より')
                yobi.append(rule[3])
                lat.append(latex(Integral(rule[3],rule[4]))+'='+latex(integrate(rule[3],rule[4])))
                lat.append('となる. ')
        elif (rule_kind.replace('ConstantRule','')!=rule_kind):
                if (len(lat)==0):
                    lat.append('積分の公式より')
                lat.append(latex(Integral(rule[1],rule[2]))+'='+latex(rule[1]*rule[2]))
                lat.append('となる. ')
        elif (rule_kind.replace('DontKnowRule','')!=rule_kind):
                if (((str(rule).replace('sin('+str(moji)+')','')!=str(rule))or(str(rule).replace('cos('+str(moji)+')','')!=str(rule))or(str(rule).replace('tan('+str(moji)+')','')!=str(rule)))and(str(rule).replace('/','')!=str(rule))):
                    wei=rule[0].subs(simplify('sin('+latex(rule[1])+')'),simplify('2*t/(1+t**2)'))
                    wei=wei.subs(simplify('cos('+latex(rule[1])+')'),simplify('(1-t**2)/(1+t**2)'))
                    wei=wei.subs(simplify('tan('+latex(rule[1])+')'),simplify('(2*t)/(1-t**2)'))
                    wei=wei*(simplify('2/(1+t**2)'))
                    wei=wei.subs(moji,2*atan(t))
                    if (str(integrate(wei,simplify('t')))==str(integrate(wei,simplify('t'))).replace('Integ','')):
                        weie='　t=\\tan{\\left(\\frac{'+str(moji)+'}{2} \\right)}とおくと, \\sin{\\left('+str(moji)+' \\right)}=\\frac{2 t}{1+t^{2}},\\cos{\\left('+str(moji)+' \\right)}=\\frac{1-t^{2}}{1+t^{2}},\\tan{\\left('+str(moji)+' \\right)}=\\frac{2 t}{1-t^{2}},d'+str(moji)+'=\\frac{2}{t^{2} + 1}dtであるから(ワイエルシュトラス置換)'
                        lat=ent(lat,weie)
                        lat.append(latex(Integral(rule[0],rule[1]))+'='+latex(Integral(wei,simplify('t'))))
                        lat.append('となり,')
                        rec_w=integrate(wei,simplify('t')).subs(simplify('t'),simplify('tan('+str(moji)+'/2)'))
                        lat,yobi=steps(integral_steps(wei,simplify('t')),lat,yobi,moji)
                        lat[-1]=lat[-1]+'したがって'
                        lat.append(latex(Integral(wei,simplify('t')))+'='+latex(integrate(wei,simplify('t')))+'='+latex(rec_w))
                        lat.append('となる. 以上より')
                        lat.append(latex(Integral(rule[0],rule[1]))+'='+latex(rec_w)+'+Const.')
                        sim=simplify(rec_w)
                        if (latex(sim)!=latex(rec_w)):
                            lat.append('となり, これを変形すると')
                            lat.append(latex(Integral(rule[0],rule[1]))+'='+latex(sim))
                            lat.append('となる.')
                    else:
                        if (latex(Integral(rule[0],rule[1]))==latex(integrate(rule[0],rule[1]))):
                            lat=['解析的な結果が見つかりません.',latex(Integral(rule[0],rule[1]))+'=?']
                        else:
                            if (len(lat)==0):
                                lat.append('積分の公式より')
                            lat.append(latex(Integral(rule[0],rule[1]))+'='+latex(integrate(rule[0],rule[1])))
                            lat.append('となる. ')            
                else:
                        if (latex(Integral(rule[0],rule[1]))==latex(integrate(rule[0],rule[1]))):
                            lat=['解析的な結果が見つかりません.',latex(Integral(rule[0],rule[1]))+'=?']
                        else:
                            if (len(lat)==0):
                                lat.append('積分の公式より')
                            lat.append(latex(Integral(rule[0],rule[1]))+'='+latex(integrate(rule[0],rule[1])))
                            lat.append('となる. ')    

            

        elif (rule_kind.replace('AddRule','')!=rule_kind):
                for ii in range(len(rule[0])):
                        
                        if (check_rule(str(rule[0][ii])[0:15])==1):
                            lat=steps(rule[0][ii],lat,yobi,moji)[0]
                            lat.pop(-1)
                            if(ii>0)and(len(lat)%2==1):
                                    lat[-2]=lat[-2]+', '+lat[-1]
                                    lat.pop(-1)
                            if (ii==len(rule[0])-1):
                                lat.append('となり, ')
                        else:
                            if (ii>0)and(len(lat)%2==0):
                                lat.append('となる. また, ')
                            lat=steps(rule[0][ii],lat,yobi,moji)[0]
                lat[-1]=lat[-1]+'これらの式を足し合わせると'
                yobi.append(rule[1])
                lat.append(latex(Integral(rule[1],rule[2]))+'='+latex(integrate(rule[1],rule[2])))
                lat.append('となる. ')        
        elif (rule_kind.replace('RewriteRule','')!=rule_kind):
                lat=ent(lat,'式を変形すると, ')
                lat.append(latex(rule[2])+'='+latex(rule[0]))
                yobi.append(rule[2])
                lat.append('となる. ')
                lat=(steps(rule[1],lat,yobi,moji))[0]
        elif (rule_kind.replace('URule','')!=rule_kind):
                lat=ent(lat,'　'+latex(rule[0])+'='+latex(rule[1])+'とおくと, '+'d'+latex(rule[0])+'=('+latex(diff(rule[1],rule[5]))+')d'+latex(rule[5])+'となるから')#初めに記述する文字の関数
                lat.append(latex(Integral(rule[4].subs(rule[0],rule[1]),rule[5]))+'=')
                rec=len(lat)
                rec_y=len(yobi)
                lat.append('となる. ここで,')
                lat,yobi=(steps(rule[3],lat,yobi,moji))
                inte=yobi[rec_y].replace(rule[0],rule[1])*diff(rule[1],rule[5])
                lat[rec-1]=latex(Integral(inte,rule[5]))+'='+latex(Integral(yobi[rec_y],rule[0]))
                lat[rec+1]=latex(Integral(yobi[rec_y],rule[0]))+'='+latex(integrate(yobi[rec_y],rule[0]))
                lat[-1]=lat[-1]+'よって'
                lat.append(latex(Integral(inte,rule[5]))+'='+latex(integrate(inte,rule[5])))
                lat.append('となる.')
        elif (rule_kind.replace('AlternativeRule','')!=rule_kind):
                if len(lat)==0:
                    for i in range(len(rule[0])):
                        lat.append('　Method'+str(i+1)+':   ')
                        lat,yobi=(steps(rule[0][i],lat,yobi,moji))
                        lat[-1]=lat[-1]+'以上より'
                        lat.append(latex(Integral(yobi[0]))+"="+latex(simplify(integrate(yobi[0])))+'+Const.')
                        lat.append('となる. ')
                        if len(rule[0])-i!=1:
                            lat.append('  ')
                            lat.append('  ')
                            lat.append('  ')
                else:
                    lat,yobi=(steps(rule[0][0],lat,yobi,moji))
        elif (rule_kind.replace('CyclicPartsRule','')!=rule_kind):
            subrule=[]
            subrule.append('PartsRuledel')
            subrule.append(rule[0][0][0])
            subrule.append(rule[0][0][1])
            subrule.append(rule[0][0][2])
            subrule.append('')
            subrule.append(rule[2])
            subrule.append(rule[3])
            lat,yobi=(steps(subrule,lat,yobi,moji))
            subrule=[]
            subrule.append('PartsRuledel')
            subrule.append(rule[0][1][0])
            subrule.append(rule[0][1][1])
            subrule.append(rule[0][1][2])
            subrule.append('')
            subrule.append(yobi[-1])
            subrule.append(rule[3])
            lat,yobi=(steps(subrule,lat,yobi,moji))
        elif (rule_kind.replace('PartsRule','')!=rule_kind):
            lat=ent(lat,"("+latex(integrate(rule[1],rule[5]))+")'="+latex(rule[1])+"として, 部分積分を行うと")
            lat.append(latex(Integral(rule[4],rule[5]))+'='+latex(rule[0]*integrate(rule[1],rule[5])-Integral(diff(rule[0],rule[5])*integrate(rule[1],rule[5]),rule[5])))
            yobi.append(rule[4])
            if (rule[3]!=''):
                lat.append('となる. さらに,')
                lat,yobi=(steps(rule[3],lat,yobi,moji))
            else:
                lat.append('となる.')
        return lat, yobi    
def set_step(form,ini,end,ss,lat,moji):
    x,y,z,t,s = symbols('x y z t s')
    moji=simplify(moji)
    if (ss==1):
        ans=integrate(form,(moji,simplify(ini),simplify(end)))
        if (str(ans)!=str(ans).replace('Integ',''))and(str(ans)==str(ans).replace('Piecewise','')):
            f = lambdify(moji, form, 'numpy')
            yy = lambda moji: f(moji)
            if ini=='oo':
                ini=10**10
            if end=='oo':
                end=10**10
            if ini=='-oo':
                ini=-10**10
            if end=='oo':
                end=-10**10                
            ans=quad(yy, int(ini), int(end))[0]
    t1=time.time()
    rule=integral_steps(form, moji)
    t2=time.time()
    ini=str(ini).replace('oo','\\infty')
    end=str(end).replace('oo','\\infty')  
    if (t2-t1>10)and(str(rule)[0:15].replace('DontKnowRule','')!=str(rule)[0:15]):
        lat=['解析的な結果が見つかりません.',latex(Integral(rule[0],rule[1]))+'=?']
    else:
        yobi=[form]
        lat=steps(rule,lat,yobi,moji)[0]
        if len(lat)==0:
            lat.append('積分公式より')
            lat.append(latex(Integral(form,moji))+'='+latex(integrate(form,moji))+'+Const')
            

        if ((str(rule)[0:15]==str(rule)[0:15].replace('AlternativeRule',''))and(len(lat)>3)and(str(rule)[0:15]==str(rule)[0:15].replace('DontKnowRule',''))):
            lat[-1]=lat[-1]+'以上より'
            lat.append(latex(Integral(form,moji))+'='+latex(simplify(integrate(form,moji)))+'+Const')
        if (ss==1):
            if (lat[-1]==lat[-1].replace('=?','')):
                if len(lat)%2==1:
                    lat[-1]='となるから'
                else:
                    lat.append('となるから')
                if str(ans)==str(simplify(ans)):
                    lat.append('\\int_{'+ini+'}^{'+end+'}('+latex(form)+')dx=\\left['+latex(integrate(form,moji))+'\\right]_{'+ini+'}^{'+end+'}='+latex(ans))
                else:
                    lat.append('\\int_{'+ini+'}^{'+end+'}('+latex(form)+')dx=\\left['+latex(integrate(form,moji))+'\\right]_{'+ini+'}^{'+end+'}='+latex(ans)+'='+latex(simplify(ans)))
            else:
                lat.append('近似解:')
                lat.append('\\int_{'+ini+'}^{'+end+'}('+latex(form)+')dx\\approx'+latex(ans))

    return lat
def ent(lat,eq):
    if (len(lat)==0):
        lat.append(eq)
    elif (lat[-1]==' '):
        lat.append(eq)
    else:
        lat[-1]=lat[-1]+eq
    return lat
def check_rule(rule):
    che=0
    lis=['DontKnowRule','ConstantRule','Fresnel','ErfRule', 'PiecewiseRule','ReciprocalRule','CiRule','SiRule','LiRule','TrigRule','Power','Exp']
    for i in range(len(lis)):
        if (rule.replace(lis[i],'')!=rule):
            che=1
    return che



def dif(a,lat):
    a=str(a)
    xi=symbols('xi')
    count=0
    rec=[]
    aa=str(a)
    ini=-1
    for i in range(len(str(a))):
        if str(a)[i]=='(':
            if count==0:
                ini=i
            count=count+1
        elif str(a)[i]==')':
            count=count-1
            if count==0:
                
                rec.append(str(a)[ini+1:i])
                aa=aa.replace('('+str(a)[ini+1:i]+')','(xi)',1)
    if (ini!=-1)and(len(lat)<20):
        for i in range(10):
            aa=aa.replace(str(i)+'*',str(i)+'$')
            aa=aa.replace('*'+str(i),'$'+str(i))
        aa=aa.replace('*','#')
        kou=re.split('[-+]',aa)
        p=0#置換の数を調整
        for i in range(len(kou)):

            if len(re.split('#',kou[i],1))>1:
                seki=re.split('#',kou[i],1)
                kou[i]=kou[i].replace('#','*')
                seki[0]=seki[0].replace('#','*')
                seki[1]=seki[1].replace('#','*')
                seki[0]=seki[0].replace('$','*')
                seki[1]=seki[1].replace('$','*')
                if (seki[0]!=seki[0].replace('xi','')):
                    seki[0]=seki[0].replace('xi',rec[i+p],1)
                    p=p+1
                for j in range(20):
                    if (seki[1]!=seki[1].replace('xi','')):
                        seki[1]=seki[1].replace('xi',rec[i+p],1)
                        p=p+1
                    else:
                        break                   
                lat.append('積の微分公式より')
                lat.append('('+latex(simplify(seki[0])*simplify(seki[1]))+")'=("+latex(simplify(seki[0]))+")'("+latex(simplify(seki[1]))+')+('+latex(simplify(seki[0]))+')('+latex(simplify(seki[1]))+")'")
                lat.append('となり,')
                lat=dif(seki[0],lat)
                lat=dif(seki[1],lat)
            else:
                kou[i]=kou[i].replace('$','*')
                kou[i]=kou[i].replace('#','*')
                if re.sub('[1-9]','',rec[i+p])!='':
                    ko=simplify(diff(simplify(kou[i]),xi))
                    if (ko!=0)and(ko!=ko.replace(xi,x)):
                        if (rec[i+p]!=rec[i+p].replace('{','')):
                            lat.append("("+latex(simplify(kou[i].replace('xi',rec[i+p])))+")'="+latex(simplify(str(ko).replace('xi',rec[i+p])))+latex(simplify(rec[i+p]))+"'")
                            lat.append('となる.')
                            lat=dif(rec[i+p],lat)
                        else:
                            lat.append("("+latex(simplify(kou[i].replace('xi',rec[i+p])))+")'=("+latex(simplify(str(ko).replace('xi',rec[i+p])))+")("+latex(diff(simplify(rec[i+p]),x))+")")
                            lat.append('となる.')   
                    else:
                        lat.append(latex(ko).replace('xi',rec[i+p]))
                else:
                    lat.append(latex(simplify(kou[i].replace('xi',rec[i+p]))))
    if len(lat)==0:
        lat.append('微分の公式より')
        lat.append('('+latex(simplify(a))+")'="+latex(diff(simplify(a),x)))
        lat.append('となる.')
    return lat

def mat(aaa):
    lis=['sin','cos','tan','sqrt','log','exp']
    aaa=re.sub('{',"(",aaa)
    aaa=re.sub('}',")",aaa)
    aaa=aaa.replace('^','**')
    aaa=aaa.replace("\\","")
    count=0
    ini=0
    for j in range(10):
        for i in range(len(str(aaa))):
            if (str(aaa)[i-5:i]=='frac(')and(ini==0):
                aaa=aaa[:i-5]+'(    '+aaa[i:]
                ini=1
                count=1
            if (str(aaa)[i]=='(')and(ini==1):
                count=count+1
            elif (str(aaa)[i]==')')and(ini==1):

                count=count-1
                if count==0:
                    aaa=aaa[:i+1]+'/'+aaa[i+1:]
                    ini=0
        if aaa==aaa.replace('frac(',''):
            break
    aaa=re.sub('    |right|left',"",aaa)
    for i in range(9):
        for j in range(len(lis)):
            aaa=re.sub(str(i)+' '+lis[j],str(i)+'*'+lis[j],aaa)
    for i in range(len(lis)):
        aaa=aaa.replace(') '+lis[i],')*'+lis[i])
    return simplify(aaa)

def maxi(form,check2):
    decimal.getcontext().prec = 6
    lat=[]
    form2=latex(form)
    form=str(form)
    lat.append('問.')
    lat.append(latex(simplify(form))+'の極値を求めよ.')
    lat.append('解答. 関数f(x)='+latex(simplify(form))+'を微分すると')
    dd=diff(simplify(form),x)
    if (simplify(dd)!=dd):
        lat.append("f'(x)="+latex(dd)+'='+latex(simplify(dd)))
    else:
        lat.append("f'(x)="+latex(dd))
    res,su,lat=equa(simplify(dd),lat)
    br=0
    for i in range(len(res)):
        if (i<len(res)):
            if sympy.im(simplify(form).replace(x,N(res[i-br])))!=0:
                lat[-1]=lat[-1]+'(x='+latex(res[i])+','
                res.pop(i)

                br=br+1
    if br>0:
        lat[-1]=lat[-1][:-1]+'は定義域外)'
    if len(res)==0:
        try:
            if (dd.replace(x,1)>0):
                lat.append("となるため, f'(x)=0は解を持たず, f'(x)は常に正であるから, 極小値と極大値は存在しない.")
            else:
                lat.append("となるため, f'(x)=0は解を持たず, f'(x)は常に負であるから, 極小値と極大値は存在しない.")
        except:
            lat.append("となるため, f'(x)=0は解を持たず, 極小値と極大値は存在しない.")

    else:
        lat.append('となる. よって, 増減表は次のようになる.')
        zou='\\begin{array}{c|'+'ccccccccccccccccc'[:2*len(res)+1]+'}x & \\cdots'
        for i in range(len(res)):
            zou=zou+'& '+latex(res[i])+' & \\cdots'
        zou=zou+'\\\\ \\hline f’(x) '
        for i in range(len(res)):
            if (dd.replace(x,sympy.re(N(res[i])-10**(-5)))>0):
                zou=zou+'& '+'+'+' & 0'
            else:    
                zou=zou+'& '+'-'+' & 0'
        if (dd.replace(x,sympy.re(N(res[-1])+10**(-5)))>0):
            zou=zou+'& + \\\\ \\hline f(x)  '
        else:
            zou=zou+'& - \\\\ \\hline f(x)  '
        kyo=[]
        for i in range(len(res)):
            if (dd.replace(x,sympy.re(N(res[i])-10**(-5)))>0):
                kyo.append(1)
                if (su==0):
                    try:
                        zou=zou+'& '+'\\nearrow'+' & '+ latex(simplify(simplify(form).replace(x,res[i])))
                    except:
                        zou=zou+'& '+'\\nearrow'+' & '+ latex(simplify(simplify(form).replace(x,float(decimal.Decimal(float(N(res[i])))+0))))

                else:
                    zou=zou+'& '+'\\nearrow'+' & '+ latex(simplify(simplify(form).replace(x,float(decimal.Decimal(float(res[i]))+0))))
            else:
                kyo.append(-1)
                if (su==0):
                    try:
                        zou=zou+'& '+'\\searrow'+' & '+ latex(simplify(simplify(form).replace(x,res[i])))
                    except:
                        zou=zou+'& '+'\\searrow'+' & '+ latex(simplify(simplify(form).replace(x,float(decimal.Decimal(float(N(res[i])))+0))))

                else:
                    zou=zou+'& '+'\\searrow'+' & '+ latex(simplify(simplify(form).replace(x,float(decimal.Decimal(float(res[i]))+0))))

        if (dd.replace(x,sympy.re(N(res[-1])+10**(-5)))>0):
            zou=zou+'& '+'\\nearrow'+' & \\end{array}'
            kyo.append(1)
        else:
            kyo.append(-1)
            zou=zou+'& '+'\\searrow'+' & \\end{array}'      
        lat.append(zou)
        lat.append('以上より,この関数は')
        ky='極大値:'
        for i in range(len(res)):
            if (kyo[i]>0)and(kyo[i+1]<0):
                try:
                    kk=simplify(simplify(form).replace(x,res[i]))
                    err=latex(kk)
                except:
                    kk=simplify(simplify(form).replace(x,float(N(res[i]))))

                if (su==0):
                    ky=ky+'x='+latex(res[i])+'\\approx'+latex(sympy.re(N(res[i])))+'のとき'+latex(kk)+'\\approx'+latex(sympy.re(N(kk)))+'\\\\　'
                else:
                    ky=ky+'x='+latex(float(decimal.Decimal(res[i])+0))+'のとき'+latex(float(decimal.Decimal(float(N(kk)))+0))+'\\\\　'
 
        if (ky=='極大値:'):
            ky='極大値:なし,'
        ky=ky[:-1]
        lat.append(ky)
        lat.append('となり, ')
        
        ky='極小値:'

        for i in range(len(res)):
            if (kyo[i]<0)and(kyo[i+1]>0):
                try:
                    kk=simplify(simplify(form).replace(x,res[i]))
                    err=latex(kk)
                except:
                    kk=simplify(simplify(form).replace(x,float(N(res[i]))))
                if (su==0):
                    ky=ky+'x='+latex(res[i])+'\\approx'+latex(sympy.re(N(res[i])))+'のとき'+latex(kk)+'\\approx'+latex(sympy.re(N(kk)))+'\\\\　'
                else:
                    ky=ky+'x='+latex(float(decimal.Decimal(res[i])+0))+'のとき'+latex(float(decimal.Decimal(float(N(kk)))+0))+'\\\\　'
        if (ky=='極小値:'):
            ky='極小値:なし,'
        ky=ky[:-1]
        lat.append(ky)
        lat.append('となる. ')
        if check2!='yes':
            lat=['　'+form2+'の極値',lat[-4],'',lat[-2]]
    return lat
def solvenew(expr,ma):

    x = symbols('x')
    f = lambdify(x, expr, 'numpy')
    res=[]
    h = 1E-10
    res=[]
    ch=0
    for j in range(20000):
        if len(res)>ma:
            break
        try:
            x=(j/200)*(-1)**j
            if (type(f(x))!=complex)and(type(f(x-1/100*(-1)**j))!=complex):
                if (f(x)*f(x-1/100*(-1)**j)<=0):
                    ch=ch+1
                    ty=abs(f(x)-f(x-1/100*(-1)**j))
                    if (ty==0):
                        ty=10**(-10)
                    for i in range(1000):
                        try:
                            x = x - float(f(x)*2*h/(f(x + h)-f(x - h)))
                        except:
                            
                            x=x+10*h
                        
                        if abs(f(x)) < 10**(-10)*ty:
                            
                            if (len(res)>0):
                                tyo=0
                                for k in range(len(res)):
                                    if (abs(res[k]-x)<10**(-5)):
                                        tyo=1
                                if tyo==0:
                                    if (f(x-10**(-5))*f(x+10**(-5))<0)or(abs(f(x)/f(x+10**(-5)))<10**(-5)):
                                        if (abs(x)<10**(-8)):
                                            res.append(0)
                                        else:
                                            res.append(x)
                            else:
                                if (f(x-10**(-5))*f(x+10**(-5))<0)or(abs(f(x)/f(x+10**(-5)))<10**(-5)):
                                    if (abs(x)<10**(-8)):
                                        res.append(0)
                                    else:
                                        res.append(x)
                            break
        except:
            x=0
    if (ch<3):
        for j in range(100):
            x=(j)*(-1)**j
            try:
                for i in range(100):
                        try:
                            x = x - float(f(x)*2*h/(f(x + h)-f(x - h)))
                        except:
                            x=x+100*h
                        if abs(f(x)) <= 10**(-25):
                            if (len(res)>0):
                                tyo=0
                                for k in range(len(res)):
                                    if (abs(res[k]-x)<10**(-5)):
                                        tyo=1
                                if tyo==0:
                                    if (f(x-10**(-5))*f(x+10**(-5))<0)or(abs(f(x)/f(x+10**(-5)))<10**(-5)):
                                        if (abs(x)<10**(-8)):
                                            res.append(0)
                                        else:
                                            res.append(x)
                            else:
                                if (f(x-10**(-5))*f(x+10**(-5))<0)or(abs(f(x)/f(x+10**(-5)))<10**(-5)):
                                    if (abs(x)<10**(-8)):
                                        res.append(0)
                                    else:
                                        res.append(x)
                            break
            except:
                x=0
    res.sort()
    return res


def det(m,lat):
    lat.append(latex(sympy.Matrix(m)))
    lat.append('')
    for i in range(4):
        if i%2==0:
            lat[-1]=lat[-1]+'('+latex(m[0][i])+')'+latex(sympy.Matrix(minor_submatrix(m,0,i)))+'-'
        else:
            lat[-1]=lat[-1]+'('+latex(m[0][i])+')'+latex(sympy.Matrix(minor_submatrix(m,0,i)))+'+'
    lat.append('')
    for i in range(4):
        mm=sympy.Matrix(minor_submatrix(m,0,i))
        lat[-1]=lat[-1]+'det'+latex(mm)+'='+latex(sympy.det(mm))+','
    return lat


def minor_submatrix(m, i, j):
   mtrx = []
   for row in range(len(m)):
       if row != i:
           mtrx_rows = []
           for col in range(len(m[0])):
               if col != j:
                   mtrx_rows.append(m[row][col])
           mtrx.append(mtrx_rows)
   return mtrx

def equa(form,r):
    try:
        su=0
        res=solve(simplify(form),x)
        if len(res)==0:
            su=1/0
    except:
        su=1
        res=solvenew(simplify(form),9)
    try:   
        bre=0
        for j in range(len(res)):
            for i in range(len(res)-1):
                if (sympy.re(N(res[i]))>sympy.re(N(res[i+1]))):
                    ch=res[i+1]
                    res[i+1]=res[i]
                    res[i]=ch  

        for j in range(10):
            for i in range(len(res)):
                if (i<len(res)):
                    if abs(sympy.im(N(res[i])))>10**(-15):
                        
                        res.pop(i)
                        bre=bre+1 
            if (j>bre):
                break
        bre=0
        for j in range(10):  
            for i in range(len(res)):
                if (i<len(res))and(i>0):
                    if (sympy.re(res[i-1])==sympy.re(res[i])):
                        res.pop(i)
                        bre=bre+1                    
                if (i<len(res)):
                    if (str(res[i])!=str(res[i]).replace('CR','')):
                        res[i]=sympy.re(N(res[i]))
            if (j>bre):
                break
        if (r!=''):
            fa=factor(simplify(form))
            r.append('　'+latex(form)+'=0について解く.')
            if (form!=simplify(form)):
                r[-1]=r[-1]+'この式を変形して'
                r.append(latex(simplify(form))+'=0')
                r.append('となり,')
            tr=0
            if (len(fa.args)>0):
                if (ask(Q.rational(fa.args[0]))==True)and(len(fa.args)>2):
                    tr=1
                elif (ask(Q.irrational(fa.args[0]))==True)and(len(fa.args)>2):
                    tr=1
                elif(len(fa.args)>1):
                    tr=1
            if (str(fa)!=str(expand(form)))and(tr==1)and((str(srepr(fa))[:4]!=str(srepr(fa))[:4].replace('Mul',''))or(str(srepr(fa))[:4]!=str(srepr(fa))[:4].replace('Pow',''))):

                r[-1]=r[-1]+'この式を因数分解すると'
                r.append(latex(fa))
                r.append('となり')
                if (len(res)!=0):
                    r.append('x=')
                    for i in range(len(res)):
                        r[-1]=r[-1]+latex(res[i])+','
                    r[-1]=r[-1][:-1]
                else:
                    r.append("\\{x\\in \mathbb{R}|"+latex(form)+"=0\\}=\\varnothing")
            else:
                try:
                    po=Poly(form,x)
                    a,b,c,d,e=symbols('a b c d e')
                    if len(po.all_coeffs())>5:
                        r[-1]=r[-1]+'この方程式の次数は5次以上であるから, アーベル-ルフィニの定理より, 解の公式は存在しない.'
                    elif (len(po.all_coeffs())==2):
                        r[-1]=r[-1]+'　ax+b=0の解の公式は'+latex(x)+'='+latex(solve(a*x+b,x)[0])+'となるから,a='+latex(po.all_coeffs()[0])+',b='+latex(po.all_coeffs()[1])+'を代入して'
                    elif (len(po.all_coeffs())==3):
                        r[-1]=r[-1]+'　ax^{2}+bx+c=0の解の公式は'
                        kai=solve(a*x**2+b*x+c,x)
                        r.append('x_{1}\\: &='+latex(kai[0])+'\\\\x_{2}&='+latex(kai[1]))
                        r.append('となり,a='+latex(po.all_coeffs()[0])+',b='+latex(po.all_coeffs()[1])+',c='+latex(po.all_coeffs()[2])+'を代入して')
                    elif (len(po.all_coeffs())==4):
                        r[-1]=r[-1]+'　ax^{3}+bx^{2}+cx+d=0の解の公式はカルダノの公式より'
                        kai=solve(a*x**3+b*x**2+c*x+d,x)
                        r.append('x_{1}\\: &='+latex(kai[0])+'\\\\x_{2}&='+latex(kai[1])+'\\\\x_{3}&='+latex(kai[2]))
                        r.append('となり,a='+latex(po.all_coeffs()[0])+',b='+latex(po.all_coeffs()[1])+',c='+latex(po.all_coeffs()[2])+',d='+latex(po.all_coeffs()[3])+'を代入して')
                        
                    elif (len(po.all_coeffs())==5):
                        r[-1]=r[-1]+'解の公式はフェラーリの公式より'
                        r.append(latex(x)+'='+latex(solve(a*x**4+b*x**3+c*x**2+d*x+e,x)[0]))
                        r.append('となるから')     
                    if (len(res)!=0):
                        r.append('x=')
                        for i in range(len(res)):
                            r[-1]=r[-1]+latex(res[i])+','
                        r[-1]=r[-1][:-1]
                    else:
                        r.append("\\{x\\in \mathbb{R}|"+latex(form)+"=0\\}=\\varnothing")                                 
                except PolynomialError:
                    if (len(res)!=0):
                        r.append('x=')
                        for i in range(len(res)):
                            r[-1]=r[-1]+latex(res[i])+','
                        r[-1]=r[-1][:-1]
                    else:
                        r.append("\\{x\\in \mathbb{R}|"+latex(form)+"=0\\}=\\varnothing")
        else:
            r=[]
            r.append('　'+latex(form)+'=0の解は')
            if (len(res)!=0):
                r.append('x=')
                for i in range(len(res)):
                    r[-1]=r[-1]+latex(res[i])+','
                r[-1]=r[-1][:-1]
            else:
                r.append("\\{x\\in \mathbb{R}|"+latex(form)+"=0\\}=\\varnothing")

    except TypeError:
        su=2
    return res,su,r

def lap(form,ss):
    lat=[]
    s=Symbol('s',irrational=True,positive=True)
    try:
        la=laplace_transform(form, x, s)[0]
        
        if (ss==1):
            lat.append('ラプラス変換は次の式を満たす.')
            lat.append('\\mathcal{L}_{x} \\left[ f(x) \\right](s)  = \\int_{0}^{\\infty} f(x) e^{-sx} dx')
            lat.append('ここで, f(x)='+latex(form)+'として, \\int_{0}^{\\infty}('+latex(form)+')e^{-sx} dxを計算する.')
            lat=set_step(form*exp(-s*x),0,oo,1,lat,'x')
            lat.append('よって, '+latex(form)+'をラプラス変換すると')
        else:
            lat.append('')
        lat.append('\\mathcal{L}_{x}\\left['+latex(form)+'\\right](s)='+latex(la))
        lat.append('となる.')
    except:
        lat.append('解析的な結果が見つかりません.')
        lat.append('\\mathcal{L}_{x}\\left['+latex(form)+'\\right](s)=?')
    return lat

def inlap(form,ss):
    lat=[]
    s,x=symbols('s x')
    form=simplify(form)
    if (diff(form,s)==0):
        lat.append('デルタ関数の性質より')
        lat.append('\\int_{-\\infty}^{\\infty} f(x) \\delta(x) dx=f(0)')
        lat.append('となる. よって')
        lat.append('\\mathcal{L}_{x} \\left[ \\delta(x) \\right](s)  = \\int_{0}^{\\infty} \\delta(x) e^{-sx} dx=e^{-0}=1')
        if form==1:
            lat.append('となる. よって')
        else:
            lat.append('となるから,\\mathcal{L}^{-1}_{s} \\left[a F(s)\\right](x)=af(x)を利用して')
        ss=2
    if (ss!=2):
        lat.append('逆ラプラス変換において, 次の式を満たす(ブロムウィッチ積分).')
        lat.append('\\mathcal{L}^{-1}_{s} \\left[ F(s) \\right](x)  = \\frac{1}{2 \\pi i}\\int_{c-i\\infty}^{c+i\\infty} F(s) e^{sx} ds')
        lat.append('ここで,F(s)e^{sx}の留数をR_{1},R_{2},\\cdots,R_{n}とすると,留数定理より')
        lat.append('\\frac{1}{2 \\pi i}\\int_{c-i\\infty}^{c+i\\infty} F(s) e^{sx} ds=R_{1}+R_{2}+\\cdots+R_{n}')
        lat.append('が従う.'+latex(form*exp(s*x))+'の特異点は')
        si=list(sympy.singularities(exp(x*s)*form,s))
        if (len(si)!=0):
            for i in range(len(si)):
                lat[-1]=lat[-1]+'s='+latex(si[i])+','
            lat[-1]=lat[-1][:-1]
            lat[-1]=lat[-1]+'となるから, それぞれの特異点における留数を求めると'
            kai=0
            lat.append('')
            for i in range(len(si)):
                res=resi(exp(x*s)*form, s, si[i]) 
                lat[-1]=lat[-1]+'R_{'+str(i)+'}=Res_{s='+latex(si[i])+'}('+latex(form*exp(x*s))+')='+latex(res)+','
                kai=kai+res
            lat[-1]=lat[-1][:-1]
            lat.append('となり')
            lat.append('R_{1}+R_{2}+\\cdots+R_{n}='+latex(kai))
            if (kai!=simplify(kai)):
                lat[-1]=lat[-1]+'='+latex(simplify(kai))
            lat.append('となる. 従って')
            if ss==0:
                lat=[]
                lat.append('')
            lat.append('\\mathcal{L}^{-1}_{s} \\left[ '+latex(form)+' \\right](x)  ='+latex(simplify(kai))+'.')
        else:
            lat=[]
            lat.append('解析的な結果が見つかりません.')
            lat.append('\\mathcal{L}^{-1}_{s} \\left['+latex(form)+' \\right](x)  =?')
    return lat
'''
def lim(form,end,ss,lat):
    lim1=limit(form,x,end,'+')
    lim2=limit(form,x,end,'-')
    end=end.replace('oo','\\infty')
    if lim1==lim2:
        if ss==1:
            lat.append('')
            lat.append('\\lim_{x \\to '+str(end)+'}'+latex(form)+'='+latex(lim1))
    else:
        lat.append('右側極限と左側極限で値が異なるため,\\lim_{x \\to '+str(end)+'}'+latex(form)+'は定まりません.')
        lat.append('\\lim_{x \\to '+str(end)+'+0}'+latex(form)+'='+latex(lim1)+',\\lim_{x \\to '+str(end)+'-0}'+latex(form)+'='+latex(lim2))
    return lat
'''
def lim(form,val,ss,lat,sit,moji,deep):
    val=simplify(val)
    ar=srepr(form)
    ar_r=form.args
    if (diff(form,moji)==0):
        lat.append('\\lim_{'+latex(moji)+' \\to '+latex(val)+'}\\left('+latex(form)+'\\right)='+latex(form))
        lat.append('となる.')
        return form,lat
    elif (str(form.subs(moji,val))!='nan')and(sit!=4):
        kai=simplify(form.subs(moji,val))
        lat[-1]=lat[-1]+'　'+latex(form)+'について, xを'+latex(val)+'に近づけると'
        lat.append('\\lim_{'+latex(moji)+'\\to '+latex(val)+'}\\left('+latex(form)+'\\right)='+latex(kai))
        lat.append('となる.')
        return kai,lat
    elif (ar[:3]=='Pow'):
        if ((limit(ar_r[0],moji,val)==0)and(abs(limit(ar_r[1],moji,val))==0))or((limit(ar_r[0],moji,val)==1)and(abs(limit(ar_r[1],moji,val))==oo))or((abs(limit(ar_r[0],moji,val))==oo)and(limit(ar_r[1],moji,val)==1)):
            lat[-1]=lat[-1]+'　対数をとって考えると'
            lat.append('\\lim_{'+latex(moji)+'\\to '+latex(val)+'}\\log{'+latex(form)+'}=\\lim_{'+latex(moji)+'\\to '+latex(val)+'}'+latex(log(ar_r[0])*ar_r[1])+'')
            lat.append('となり,')
            rec=len(lat)
            kai,lat=lim(log(ar_r[0])*ar_r[1],val,ss,lat,0,moji,deep+1)
            if (len(lat)-rec>5):
                lat[-1]=lat[-1]+'したがって'
                lat.append('\\lim_{'+latex(moji)+'\\to '+latex(val)+'}\\left('+latex(form)+'\\right)='+latex(exp(kai)))
                lat.append('となる.')
            return kai,lat
        else:
            lat=lim(ar_r[0],val,ss,lat,0,moji,deep+1)[1]
            lat=lim(ar_r[1],val,ss,lat,0,moji,deep+1)[1]
            return 0,lat
    elif (ar[:3]=='Add')and(sit!=1):
        rec=len(lat)
        ar_hutei=[]
        siki2=0
        for ii in ar_r:
            ar_sub=ii.subs(moji,val)
            if abs(ar_sub)==oo:
                ar_hutei.append(ii)
            elif str(ar_sub)=='nan':
                lat=lim(ii,val,ss,lat,0,moji,deep+1)[1]
            else:
                siki2=siki2+ii
        siki=0
        for ii in ar_hutei:
            siki=siki+ii
        if len(ar_hutei)>0:
            kai,lat=lim(siki,val,ss,lat,1,moji,deep+1)
        if (len(lat)-rec>5):
            lat[-1]=lat[-1]+'よって'
            lat.append('\\lim_{'+latex(moji)+'\\to '+latex(val)+'}\\left('+latex(form)+'\\right)='+latex(kai+siki2))
            lat.append('となる.')
        return kai,lat
    elif (str(radsimp(form).subs(moji,val))!='nan')and(sit!=4):
        kai_r=radsimp(form)
        kai=kai_r.subs(moji,val)
        lat[-1]=lat[-1]+'式を変形すると' 
        lat.append(latex(form)+'='+latex(kai_r))
        lat.append('となり, x='+latex(val)+'を'+latex(kai_r)+'に代入して')
        lat.append('\\lim_{'+latex(moji)+'\\to '+latex(val)+'}\\left('+latex(kai_r)+'\\right)='+latex(kai))
        lat.append('となる.')
        return radsimp(form).subs(moji,val),lat
    else:
        form_r=radsimp(1/form)
        form_r=1/form_r
        if (str(form_r.subs(moji,val))!='nan')and(sit!=4):
            kai=form_r.subs(moji,val)
            lat[-1]=lat[-1]+'式を変形すると'
            lat.append(latex(form)+'='+latex(form_r))
            lat.append('となり, x='+latex(val)+'を'+latex(form_r)+'に代入して')
            lat.append('\\lim_{'+latex(moji)+'\\to '+latex(val)+'}\\left('+latex(form_r)+'\\right)='+latex(kai))
            lat.append('となる.')
            return kai,lat
        elif (ar[:3]=='Mul')and(sit!=3):
            ff=1
            gg=1
            for ii in ar_r:
                if limit(ii,moji,val)!=0:
                    gg=gg/ii
                else:
                    ff=ff*ii
            dif1=simplify(diff(ff,moji)/diff(gg,moji))
            dif2=simplify(diff(1/gg,moji)/diff(1/ff,moji))
            if (len(str(dif1))<len(str(dif2))):
                dif3=dif1
            else:
                dif3=dif2
            if (str(dif3.subs(moji,val))==str(dif3.subs(moji,val)).replace('Accum',''))and(deep<10):#振動
                if sit==2:
                    lat[-1]=lat[-1]+'再度, ロピタルの定理を用いて'
                else:
                    lat[-1]=lat[-1]+'ロピタルの定理を用いて'
                lat.append('\\lim_{'+latex(moji)+'\\to '+latex(val)+'}\\left('+latex(form)+'\\right)=\\lim_{'+latex(moji)+'\\to '+latex(val)+'}\\left('+latex(dif3)+'\\right)')
                lat.append('となり,')
                rec=len(lat)
                kai,lat=lim(dif3,val,ss,lat,2,moji,deep+1)
                if (len(lat)-rec>5):
                    lat[-1]=lat[-1]+'よって'
                    lat.append('\\lim_{'+latex(moji)+'\\to '+latex(val)+'}\\left('+latex(form)+'\\right)='+latex(kai))
                    lat.append('となる.')
            else:
                kai,lat=lim(form,val,ss,lat,3,moji,deep+1)
            return kai,lat    
        else:
            kai=limit(form,moji,val)
            lat[-1]=lat[-1]+'　'+latex(form)+'をx='+latex(val)+'付近で級数展開すると'
            try:
                lat.append(latex(form)+'='+latex(series(form,moji,val)))
                lat.append('となるから')
            except:
                lat[-1]='極限を計算すると'
            lat.append('\\lim_{'+latex(moji)+'\\to '+latex(val)+'}\\left('+latex(form)+'\\right)='+latex(kai))
            lat.append('となる.')

            return kai,lat
def limi(form,val,ss,lat,moji):
    ii=0
    for i in range(2):
        if ss==2:
            kai=lim(form,val,ss,lat,ii,moji,0)[0]
        elif ss==1:
            kai,lat=lim(form,val,ss,lat,ii,moji,0)
        elif ss==0:
            kai=lim(form,val,ss,lat,ii,moji,0)[0]
            lat=['']
            lat.append('\\lim_{'+latex(moji)+'\\to '+latex(val)+'}\\left('+latex(form)+'\\right)='+latex(kai)+'.')
        if (kai==0):
            try:
                k1=form.subs(moji,N(val)+10**(-10))
                k2=form.subs(moji,N(val)-10**(-10))
                if (abs(sympy.re(k1))<10**(-5))or(abs(sympy.re(N(k2)))<10**(-5)):
                    break
                else:
                    ii=4
            except:
                ii=4
        else:
            break
    return kai,lat

def resi(form,moji,z0):
    for i in range(1,10):
        form_c=form*(moji-z0)**i
        for j in range(i-1):
            form_c=diff(form_c,moji)
        res=limi(form_c,z0,2,[''],moji)[0]
        if (abs(res)!=oo)and(str(res)==str(res).replace('Accum','')):
            res=res/(gamma(i))
            break
    return res

def parse(form): #数式に使われている変数を認識して重要と考えられる順番に並べる関数, formはsympyの数式形式で入力
    ana = srepr(form)
    rr=[]
    for ii in range(100):
        if ana!=ana.replace("Symbol('",''):
            target1 = "Symbol('"
            idx1 = ana.find(target1)
            target2 = "')"
            idx2 = ana.find(target2)
            moji=ana[idx1+len(target1):idx2]
            rr.append(moji)
            ana=ana.replace("Symbol('"+moji+"')",'')
        else:
            break
    br=0
    pr=['x','t','y','s','z']
    for jj in pr:
        for ii in range(len(rr)):
            if rr[ii]==jj:
                rr[ii]=rr[br]
                rr[br]=jj
                br=br+1
    if len(rr)==0:
        rr=['x']
    for i in range(len(rr)):
        if i>0:
            form=form.replace(simplify(rr[i]),1)
    return rr,form

''' 
cd flask
set FLASK_APP=app
set FLASK_ENV=development
flask run

cd .ipynb_checkpoints
cd Flask
git init
git add .
git commit -m "update"
heroku login

git remote
git push heroku-staging main
'''



